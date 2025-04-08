import torch
from transformers import BertTokenizer, BertModel
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import os
from flask import Flask, render_template, request
import re
import random  # 导入 random 模块
import logging

# 设置日志配置
logging.basicConfig(level=logging.DEBUG, format="%(asctime)s - %(message)s")

# 加载中文预训练的 BERT 模型和 tokenizer
tokenizer = BertTokenizer.from_pretrained("bert-base-chinese")
model = BertModel.from_pretrained("bert-base-chinese")

app = Flask(__name__)

# 全局缓存：文档和对应 embedding，以及文档的第一句话
documents = {}
docs_embeddings = {}
doc_first_sentences = {}
doc_categories = {}

def get_bert_embedding(text):
    """将文本转换为 BERT 向量（句向量）"""
    inputs = tokenizer(text, return_tensors="pt", padding=True, truncation=True, max_length=512)
    with torch.no_grad():
        outputs = model(**inputs)
    return outputs.last_hidden_state[:, 0, :].squeeze().numpy()

def extract_categories(first_sentence):
    """从第一句话提取性格、人际交往能力、学习需求和服务类型"""
    match = re.match(r"([^\-]+)-([^\-]+)-([^\-]+)-([^\-]+)", first_sentence)
    if match:
        return {
            "personality": match.group(1).strip(),
            "communication": match.group(2).strip(),
            "study_needs": match.group(3).strip(),
            "service_type": match.group(4).strip()
        }
    else:
        return None

def read_word_documents(folder_path):
    """读取文件夹中的所有 Word 文档内容，并提取第一句话"""
    docs = {}
    first_sentences = {}
    global doc_categories
    doc_categories = {}

    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            doc_path = os.path.join(folder_path, filename)
            doc = Document(doc_path)
            text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
            docs[filename] = text

            # 提取第一句话
            first_sentence = text.split("\n")[0] if text else ""
            first_sentences[filename] = first_sentence

            # 提取分类信息
            categories = extract_categories(first_sentence)
            if categories:
                doc_categories[filename] = categories

    return docs, first_sentences

def preload_documents(folder_path):
    """预加载所有 Word 文档和对应的 BERT 向量"""
    global documents, docs_embeddings, doc_first_sentences
    documents, doc_first_sentences = read_word_documents(folder_path)
    docs_embeddings = {name: get_bert_embedding(text) for name, text in documents.items()}

def find_best_match(user_input, category=None, character=None, communication=None, study_needs=None):
    """找到与用户输入最匹配的文档"""
    input_embedding = get_bert_embedding(user_input)
    match_scores = []

    filtered_docs = docs_embeddings

    for doc_name, doc_embedding in filtered_docs.items():
        doc_category = doc_categories.get(doc_name, {})
        match_score = 0
        
        if category and category in doc_category.get('service_type', ''):
            match_score += 1

        if character and any(c in doc_category.get('personality', '') for c in character):
            match_score += 0.5
        
        if study_needs and any(s in doc_category.get('study_needs', '') for s in study_needs):
            match_score += 0.5

        if communication and communication in doc_category.get('communication', ''):
            match_score += 0.5
        
        if match_score > 0:
            similarity = cosine_similarity([input_embedding], [doc_embedding])[0][0]
            total_score = match_score * similarity
            match_scores.append((doc_name, total_score, match_score, similarity))
            # 打印匹配的文档和对应分数
            logging.debug(f"Document: {doc_name}, Match Score: {match_score}, Similarity: {similarity}, Total Score: {total_score}")

    match_scores = sorted(match_scores, key=lambda x: x[1], reverse=True)

    if match_scores:
        # 获取匹配得分最高的文档
        highest_score = match_scores[0][1]
        
        # 找到所有匹配得分相同的文档
        top_docs = [doc for doc in match_scores if doc[1] == highest_score]
        
        # 如果有多个文档得分相同，随机选择一个
        best_match = random.choice(top_docs)  # 随机选择
        return match_scores, best_match  # 返回所有文档和随机选择的最佳匹配
    else:
        return [], None  # 如果没有匹配项，返回空列表和 None

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        character = request.form.getlist('character')
        communication = request.form.get('communication')
        study_needs = request.form.getlist('studyNeed')
        service_type = request.form.get('serviceType')

        user_input = f"该学生性格为{'、'.join(character)}，人际交往能力为{communication}，学习方面需要{'、'.join(study_needs)}，期望的服务类型是{service_type}。"

        category = None
        for doc_name, categories in doc_categories.items():
            if service_type in categories['service_type']:
                category = categories['service_type']
                break

        match_scores, best_match = find_best_match(user_input, category, character, communication, study_needs)

        if best_match:
            matched_content = documents[best_match[0]]
            matched_content = "\n".join(matched_content.split("\n")[1:])  # 去掉文档第一行标题
            return render_template("result.html", best_match=best_match[0], score=round(best_match[1], 4),
                                   matched_content=matched_content, user_input=user_input, match_scores=match_scores)
        else:
            return render_template("result.html", best_match="没有找到匹配", score=0, matched_content="",
                                   user_input=user_input, match_scores=[])

    return render_template("index.html")

@app.route("/edit/<doc_name>", methods=["GET"])
def edit_document(doc_name):
    """显示修改页面，加载选中的文档数据"""
    best_match_data = doc_categories.get(doc_name, {})
    matched_content = documents.get(doc_name, "")
    return render_template("edit.html", best_match=doc_name, best_match_data=best_match_data, matched_content=matched_content)

@app.route("/update/<doc_name>", methods=["POST"])
def update_document(doc_name):
    """后台人员提交修改后的方案内容"""
    updated_content = request.form.get('content')

    # 更新文档的内容
    documents[doc_name] = updated_content
    docs_embeddings[doc_name] = get_bert_embedding(updated_content)

    # 保存修改后的内容到文件
    doc_path = os.path.join("database", doc_name)
    doc = Document()
    doc.add_paragraph(updated_content)
    doc.save(doc_path)

    # 提交成功后，显示成功消息并返回首页
    success_message = "提交成功！您的修改已经保存到数据库。"
    return render_template("success.html", message=success_message)

@app.route("/result/<doc_name>", methods=["GET"])
def result(doc_name):
    """展示匹配后的方案结果"""
    matched_content = documents.get(doc_name, "未找到该文档内容。")

    match_scores, best_match = find_best_match(matched_content)

    return render_template("result.html", best_match=best_match[0], score=round(best_match[1], 4),
                           matched_content=matched_content, user_input=matched_content, match_scores=match_scores)

if __name__ == "__main__":
    preload_documents("database")  # 请确保 Word 文件都在这个文件夹下
    app.run(debug=True)
